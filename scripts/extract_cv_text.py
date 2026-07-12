"""Extract plain text from CV files for analysis."""
from __future__ import annotations

import sys
from pathlib import Path

import pdfplumber
from docx import Document


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
    return "\n".join(parts)


def main(path: str) -> None:
    p = Path(path)
    if p.suffix.lower() == ".docx":
        text = extract_docx(p)
    elif p.suffix.lower() == ".pdf":
        text = extract_pdf(p)
    else:
        raise SystemExit(f"Unsupported: {p.suffix}")
    out = Path(__file__).parent / f"cv_extract_{p.stem[:30].replace(' ', '_')}.txt"
    out.write_text(text, encoding="utf-8")
    print(f"Wrote {out} ({len(text)} chars)")
    print(text[:1500])


if __name__ == "__main__":
    main(sys.argv[1])
