"""Extract plain text and optional photo from CV files."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree


def _extract_docx_paragraphs(doc: Document) -> list[str]:
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return parts


def _extract_docx_textboxes(docx_path: Path) -> list[str]:
    """Extract text from shapes/text boxes via raw XML."""
    texts: list[str] = []
    with zipfile.ZipFile(docx_path) as zf:
        xml_files = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
        for name in xml_files:
            root = etree.fromstring(zf.read(name))
            for node in root.iter():
                if node.tag.endswith("}t") and node.text and node.text.strip():
                    texts.append(node.text.strip())
    return texts


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = _extract_docx_paragraphs(doc)
    if len("\n".join(parts)) < 200:
        parts.extend(_extract_docx_textboxes(path))
    # Deduplicate consecutive duplicates while preserving order
    seen: set[str] = set()
    unique: list[str] = []
    for p in parts:
        if p not in seen:
            seen.add(p)
            unique.append(p)
    return "\n".join(unique)


def extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts)


def extract_cv_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {suffix}. Use PDF or DOCX.")


def extract_docx_first_image(path: Path) -> bytes | None:
    """Return bytes of first embedded image in DOCX, if any."""
    doc = Document(str(path))
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            return rel.target_part.blob
    with zipfile.ZipFile(path) as zf:
        media = sorted(n for n in zf.namelist() if n.startswith("word/media/"))
        if media:
            return zf.read(media[0])
    return None


def load_photo(path: Path | None) -> bytes | None:
    """Load profile photo from explicit upload only."""
    if path and path.exists():
        return path.read_bytes()
    return None
